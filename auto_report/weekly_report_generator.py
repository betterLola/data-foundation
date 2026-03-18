# -*- coding: utf-8 -*-
#!/usr/bin/env python
"""
天府市民云周报自动生成脚本

周期定义：
  本周期：上周五 ~ 本周四
  上周期：上上周五 ~ 上周四

运行方式：
  python weekly_report_generator.py
  （在每周五运行，当天即为新一周第一天，昨天（周四）数据已完整）

输出：当前目录下 周报-{年}年{月}月{d1}日-{月}月{d2}日.docx
"""
import os
import sys
import tempfile
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

import pymysql
import matplotlib
matplotlib.use('Agg')  # 非交互后端，避免弹窗
import matplotlib.pyplot as plt

from datetime import datetime, timedelta, date
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# matplotlib 中文字体
matplotlib.rcParams['font.family'] = ['SimHei', 'Microsoft YaHei', 'DejaVu Sans']
matplotlib.rcParams['axes.unicode_minus'] = False


# ─────────────────────────────────────────
# 数据库配置
# ─────────────────────────────────────────
try:
    from dotenv import load_dotenv
    load_dotenv()
except ImportError:
    pass

DB_CONFIG = {
    'host':     os.getenv('DB_HOST',     'localhost'),
    'port':     int(os.getenv('DB_PORT', '3306')),
    'user':     os.getenv('DB_USER',     'root'),
    'password': os.getenv('DB_PASSWORD', ''),
    'database': os.getenv('DB_NAME',     'daily'),
    'charset':  'utf8mb4'
}


# ─────────────────────────────────────────
# 辅助函数
# ─────────────────────────────────────────
def get_periods(anchor: date = None):
    """
    返回 (cur_start, cur_end, prev_start, prev_end)
    本周期 = 上周五 ~ 本周四
    """
    if anchor is None:
        anchor = date.today()
    days_since_thursday = (anchor.weekday() - 3) % 7
    this_thursday = anchor - timedelta(days=days_since_thursday)

    cur_end   = this_thursday
    cur_start = this_thursday - timedelta(days=6)

    prev_end   = cur_start - timedelta(days=1)
    prev_start = prev_end - timedelta(days=6)

    return cur_start, cur_end, prev_start, prev_end


def fmt_wan(n, decimals=2):
    return f"{n / 10000:.{decimals}f}万"


def fmt_date_cn(d: date):
    return f"{d.month}月{d.day}日"


def fmt_period_cn(start: date, end: date):
    if start.month == end.month:
        return f"{start.month}月{start.day}日-{end.day}日"
    return f"{start.month}月{start.day}日-{end.month}月{end.day}日"


def change_direction(pct):
    return "上升" if pct >= 0 else "下降"


def abs_pct_str(pct):
    return f"{abs(pct):.2f}%"


# ─────────────────────────────────────────
# 数据查询
# ─────────────────────────────────────────
def query_dau(cursor, start: date, end: date):
    cursor.execute("""
        SELECT
            AVG(platform_dau)       AS platform_avg,
            AVG(app_dau)            AS app_avg,
            AVG(alipay_dau)         AS alipay_avg,
            AVG(smart_frontend_dau) AS smart_avg,
            AVG(mini_program_dau)   AS mini_avg,
            MIN(platform_dau)       AS platform_min,
            MAX(platform_dau)       AS platform_max,
            COUNT(*)                AS days
        FROM platform_daily_metrics
        WHERE stat_date BETWEEN %s AND %s
    """, (start, end))
    row = cursor.fetchone()
    return {
        'platform_avg': float(row[0] or 0),
        'app_avg':      float(row[1] or 0),
        'alipay_avg':   float(row[2] or 0),
        'smart_avg':    float(row[3] or 0),
        'mini_avg':     float(row[4] or 0),
        'platform_min': int(row[5] or 0),
        'platform_max': int(row[6] or 0),
        'days':         int(row[7] or 0),
    }


def query_dau_daily(cursor, start: date, end: date):
    """返回 [(date, platform_dau), ...]"""
    cursor.execute("""
        SELECT stat_date, platform_dau
        FROM platform_daily_metrics
        WHERE stat_date BETWEEN %s AND %s
        ORDER BY stat_date
    """, (start, end))
    return cursor.fetchall()


def query_service_daily_total(cursor, start: date, end: date):
    cursor.execute("""
        SELECT DATE(stat_date), SUM(service_amount)
        FROM 5100_detail
        WHERE stat_date BETWEEN %s AND %s
        GROUP BY DATE(stat_date)
        ORDER BY DATE(stat_date)
    """, (start, end))
    return {row[0]: int(row[1]) for row in cursor.fetchall()}


def query_service_total_by_name(cursor, start: date, end: date):
    cursor.execute("""
        SELECT service_name, SUM(service_amount)
        FROM 5100_detail
        WHERE stat_date BETWEEN %s AND %s
        GROUP BY service_name
    """, (start, end))
    return {row[0]: int(row[1]) for row in cursor.fetchall()}


def query_service_peak(cursor, service_name: str, start: date, end: date):
    cursor.execute("""
        SELECT DATE(stat_date), SUM(service_amount)
        FROM 5100_detail
        WHERE service_name = %s AND stat_date BETWEEN %s AND %s
        GROUP BY DATE(stat_date)
        ORDER BY SUM(service_amount) DESC
        LIMIT 1
    """, (service_name, start, end))
    row = cursor.fetchone()
    if row:
        return row[0], int(row[1])
    return None, 0


def query_search_total_by_keyword(cursor, start: date, end: date):
    cursor.execute("""
        SELECT search_name, SUM(search_amount)
        FROM search_detail
        WHERE resource_name = 'search_behavior'
          AND stat_date BETWEEN %s AND %s
        GROUP BY search_name
    """, (start, end))
    return {row[0]: int(row[1]) for row in cursor.fetchall()}


# ─────────────────────────────────────────
# 数据分析
# ─────────────────────────────────────────
def build_platform_decline_text(cur_dau: dict, prev_dau: dict) -> str:
    """
    分析各端日活变化，生成平台备注文字。
    - 有端口下降：列出所有下降端，含均值、环比降幅、日均减少量
    - 所有端口均下降：只列降幅最大（日均减少最多）的两个端
    - 全部上升：保持原简洁格式
    """
    platforms = [
        ('APP（安卓+苹果+鸿蒙）', 'app_avg'),
        ('支付宝小程序',          'alipay_avg'),
        ('微信小程序',            'mini_avg'),
        ('智慧前端',              'smart_avg'),
    ]

    decline, rise = [], []
    for label, key in platforms:
        cur_val  = cur_dau[key]
        prev_val = prev_dau[key]
        if prev_val == 0:
            continue
        pct         = (cur_val - prev_val) / prev_val * 100
        daily_delta = cur_val - prev_val          # 负数表示下降
        if pct < 0:
            decline.append((label, cur_val, pct, daily_delta))
        else:
            rise.append((label, cur_val, pct, daily_delta))

    # 全部上升 → 原格式
    if not decline:
        return (
            f"本周期APP日活均值{fmt_wan(cur_dau['app_avg'])}（安卓+苹果+鸿蒙），"
            f"支付宝小程序日活均值{fmt_wan(cur_dau['alipay_avg'])}。"
        )

    # 所有端口均下降 → 只取日均减少量最多的两个（daily_delta 最负）
    if not rise:
        show = sorted(decline, key=lambda x: x[3])[:2]   # 最负的两个
        prefix = "本周期各端日活均有所下降，"
    else:
        show   = sorted(decline, key=lambda x: x[3])     # 全部下降端
        prefix = ""

    parts = []
    for label, cur_val, pct, daily_delta in show:
        parts.append(
            f"{label}日活均值{fmt_wan(cur_val)}，"
            f"环比下降{abs_pct_str(pct)}，"
            f"日均减少{fmt_wan(abs(daily_delta))}"
        )
    return prefix + "；".join(parts) + "。"


def analyze_dau_trend(daily_data) -> str:
    """
    根据每日 platform_dau 数据描述走势（含社治e管家+1万）。
    只描述走势形态与峰谷，不分析原因。
    """
    if not daily_data:
        return ""

    vals  = [int(v) + 10000 for _, v in daily_data]
    dates = [d for d, _ in daily_data]
    n     = len(vals)

    peak_idx   = vals.index(max(vals))
    trough_idx = vals.index(min(vals))
    peak_date  = dates[peak_idx]
    trough_date = dates[trough_idx]

    # 统计方向变化次数，判断走势类型
    direction_changes = 0
    for i in range(1, n - 1):
        prev_diff = vals[i]     - vals[i - 1]
        next_diff = vals[i + 1] - vals[i]
        if prev_diff * next_diff < 0:   # 方向翻转
            direction_changes += 1

    if direction_changes >= 2:
        trend_desc = "震荡波动"
    elif peak_idx == 0:
        trend_desc = "整体呈下降趋势"
    elif trough_idx == 0:
        trend_desc = "整体呈上升趋势"
    elif peak_idx < trough_idx:
        trend_desc = "先升后降"
    else:
        trend_desc = "先降后升"

    peak_str   = f"{fmt_date_cn(peak_date)}为{fmt_wan(max(vals))}"
    trough_str = f"{fmt_date_cn(trough_date)}为{fmt_wan(min(vals))}"

    return (
        f"{fmt_period_cn(dates[0], dates[-1])}日活{trend_desc}，"
        f"峰值出现在{peak_str}，谷值出现在{trough_str}。"
    )


def generate_dau_chart(
    cur_daily, prev_daily,
    cur_start: date, cur_end: date,
    prev_start: date, prev_end: date,
) -> str:
    """
    生成本周期与上周期日活折线对比图（含社治e管家+1万）。
    返回保存的临时图片路径。
    """
    cur_vals  = [(int(v) + 10000) / 10000 for _, v in cur_daily]
    prev_vals = [(int(v) + 10000) / 10000 for _, v in prev_daily]
    x         = list(range(len(cur_vals)))
    x_labels  = [fmt_date_cn(d) for d, _ in cur_daily]

    fig, ax = plt.subplots(figsize=(9, 4))

    cur_label  = f'本周期（{fmt_period_cn(cur_start,  cur_end)}）'
    prev_label = f'上周期（{fmt_period_cn(prev_start, prev_end)}）'

    ax.plot(x, cur_vals,  marker='o', label=cur_label,
            color='#2196F3', linewidth=2.2, markersize=6)
    ax.plot(x, prev_vals, marker='s', label=prev_label,
            color='#FF9800', linewidth=2.2, markersize=6, linestyle='--')

    # 标注本周期每日数值
    for i, v in enumerate(cur_vals):
        ax.annotate(f'{v:.2f}', (x[i], v),
                    textcoords='offset points', xytext=(0, 9),
                    ha='center', fontsize=8, color='#1565C0')

    ax.set_xticks(x)
    ax.set_xticklabels(x_labels, fontsize=9)
    ax.set_ylabel('日活（万）', fontsize=10)
    ax.set_title('日活走势对比（本周期 vs 上周期）', fontsize=12, pad=12)
    ax.legend(loc='best', fontsize=9)
    ax.grid(True, alpha=0.3, linestyle='--')
    ax.yaxis.set_major_formatter(
        plt.FuncFormatter(lambda val, _: f'{val:.1f}')
    )

    plt.tight_layout()
    tmp_path = os.path.join(tempfile.gettempdir(), 'dau_weekly_chart.png')
    plt.savefig(tmp_path, dpi=150, bbox_inches='tight', facecolor='white')
    plt.close()
    return tmp_path


def build_service_change_list(cur_dict: dict, prev_dict: dict, min_count=1000):
    all_names = set(cur_dict.keys()) | set(prev_dict.keys())
    result = []
    for name in all_names:
        cur  = cur_dict.get(name, 0)
        prev = prev_dict.get(name, 0)
        if cur < min_count and prev < min_count:
            continue
        if prev == 0:
            pct = float('inf') if cur > 0 else 0.0
        else:
            pct = (cur - prev) / prev * 100
        result.append((name, cur, prev, pct))
    result.sort(key=lambda x: x[3], reverse=True)
    return result


def build_search_rank(cur_dict: dict, prev_dict: dict, top_n=10):
    all_names = set(cur_dict.keys()) | set(prev_dict.keys())
    delta_list = []
    for name in all_names:
        cur   = cur_dict.get(name, 0)
        prev  = prev_dict.get(name, 0)
        delta = cur - prev
        delta_list.append((name, cur, prev, delta))

    delta_list.sort(key=lambda x: x[3], reverse=True)
    top_delta = delta_list[:top_n]

    cur_list = sorted(cur_dict.items(), key=lambda x: x[1], reverse=True)
    top_cur  = cur_list[:top_n]

    return top_delta, top_cur


# ─────────────────────────────────────────
# Word 文档辅助
# ─────────────────────────────────────────
def set_font(run, bold=False, size=10.5, color=None):
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = '仿宋'
    run.font.color.rgb = color if color else RGBColor(0, 0, 0)
    r    = run._r
    rPr  = r.find(qn('w:rPr'))
    if rPr is None:
        rPr = OxmlElement('w:rPr')
        r.insert(0, rPr)
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), '仿宋')
    rFonts.set(qn('w:ascii'),    '仿宋')
    rFonts.set(qn('w:hAnsi'),    '仿宋')


def add_para(doc, text='', bold=False, size=10.5, first_line_indent=True,
             left_indent=0, space_before=0, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    if first_line_indent:
        p.paragraph_format.first_line_indent = Pt(size * 2)
    if left_indent:
        p.paragraph_format.left_indent = Pt(left_indent)
    if text:
        run = p.add_run(text)
        set_font(run, bold=bold, size=size)
    return p


def add_heading_para(doc, text, size=12):
    p = doc.add_paragraph()
    p.paragraph_format.space_before      = Pt(6)
    p.paragraph_format.space_after       = Pt(4)
    p.paragraph_format.first_line_indent = Pt(0)
    run = p.add_run(text)
    set_font(run, bold=True, size=size)
    return p


def add_mixed_para(doc, segments, first_line_indent=True, left_indent=0,
                   space_before=0, space_after=6, size=10.5):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    if first_line_indent:
        p.paragraph_format.first_line_indent = Pt(size * 2)
    if left_indent:
        p.paragraph_format.left_indent = Pt(left_indent)
    for text, bold in segments:
        run = p.add_run(text)
        set_font(run, bold=bold, size=size)
    return p


def add_blank_para(doc, note='（请人工填写）', size=10.5):
    p = doc.add_paragraph()
    p.paragraph_format.space_before      = Pt(0)
    p.paragraph_format.space_after       = Pt(6)
    p.paragraph_format.first_line_indent = Pt(size * 2)
    run = p.add_run(note)
    set_font(run, bold=False, size=size, color=RGBColor(0x80, 0x80, 0x80))
    return p


def add_search_table(doc, top_delta, top_cur, size=10):
    n      = max(len(top_delta), len(top_cur))
    table  = doc.add_table(rows=n + 1, cols=4)
    table.style = 'Table Grid'

    headers = ['增量前十服务名称', '搜索次数增长量', '本周前十服务名称', '搜索次数']
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        run = cell.paragraphs[0].add_run(h)
        set_font(run, bold=True, size=size)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    for row_idx in range(n):
        row = table.rows[row_idx + 1]
        if row_idx < len(top_delta):
            name, cur, prev, delta = top_delta[row_idx]
            row.cells[0].text = ''
            row.cells[1].text = ''
            set_font(row.cells[0].paragraphs[0].add_run(name), size=size)
            r1 = row.cells[1].paragraphs[0].add_run(str(delta))
            set_font(r1, size=size)
            row.cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        if row_idx < len(top_cur):
            kw, cnt = top_cur[row_idx]
            row.cells[2].text = ''
            row.cells[3].text = ''
            set_font(row.cells[2].paragraphs[0].add_run(kw), size=size)
            r3 = row.cells[3].paragraphs[0].add_run(str(cnt))
            set_font(r3, size=size)
            row.cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

    return table


# ─────────────────────────────────────────
# 主生成逻辑
# ─────────────────────────────────────────
def generate_report(anchor: date = None):
    cur_start, cur_end, prev_start, prev_end = get_periods(anchor)

    print(f"本周期：{cur_start} ~ {cur_end}")
    print(f"上周期：{prev_start} ~ {prev_end}")

    conn   = pymysql.connect(**DB_CONFIG)
    cursor = conn.cursor()

    # ── 第一部分：日活 ────────────────────
    cur_dau        = query_dau(cursor, cur_start, cur_end)
    prev_dau       = query_dau(cursor, prev_start, prev_end)
    cur_daily_dau  = query_dau_daily(cursor, cur_start,  cur_end)
    prev_daily_dau = query_dau_daily(cursor, prev_start, prev_end)

    platform_pct = (
        (cur_dau['platform_avg'] - prev_dau['platform_avg'])
        / prev_dau['platform_avg'] * 100
        if prev_dau['platform_avg'] > 0 else 0.0
    )

    # ── 第二部分：服务人次 ────────────────
    cur_svc_daily  = query_service_daily_total(cursor, cur_start, cur_end)
    prev_svc_daily = query_service_daily_total(cursor, prev_start, prev_end)

    cur_svc_avg  = sum(cur_svc_daily.values())  / len(cur_svc_daily)  if cur_svc_daily  else 0
    prev_svc_avg = sum(prev_svc_daily.values()) / len(prev_svc_daily) if prev_svc_daily else 0
    svc_pct = (
        (cur_svc_avg - prev_svc_avg) / prev_svc_avg * 100
        if prev_svc_avg > 0 else 0.0
    )

    cur_svc_by_name  = query_service_total_by_name(cursor, cur_start, cur_end)
    prev_svc_by_name = query_service_total_by_name(cursor, prev_start, prev_end)
    change_list      = build_service_change_list(cur_svc_by_name, prev_svc_by_name)

    up_list   = [(n, c, p, pct) for n, c, p, pct in change_list if pct > 0][:3]
    down_list = [(n, c, p, pct) for n, c, p, pct in reversed(change_list) if pct < 0][:3]

    # ── 第三部分：搜索词 ─────────────────
    cur_search  = query_search_total_by_keyword(cursor, cur_start, cur_end)
    prev_search = query_search_total_by_keyword(cursor, prev_start, prev_end)
    search_ok   = bool(cur_search)
    top_delta, top_cur = (
        build_search_rank(cur_search, prev_search) if search_ok else ([], [])
    )

    cursor.close()
    conn.close()

    # ── 组装文字 ──────────────────────────
    cur_period_str  = fmt_period_cn(cur_start, cur_end)
    prev_period_str = fmt_period_cn(prev_start, prev_end)

    total_avg_str    = fmt_wan(cur_dau['platform_avg'] + 10000)
    platform_avg_str = fmt_wan(cur_dau['platform_avg'])

    dau_text_main = (
        f"{cur_period_str}：日活均值{total_avg_str}"
        f"（{platform_avg_str}+社区保障e管家约1万日活），"
        f"环比上周期（{prev_period_str}）不含社区保障e管家"
        f"{change_direction(platform_pct)}{abs_pct_str(platform_pct)}。"
    )

    svc_text_main = (
        f"本周期（{cur_period_str}）所有服务事项人次日均{fmt_wan(cur_svc_avg)}，"
        f"环比上周期（{prev_period_str}）的{fmt_wan(prev_svc_avg)}"
        f"{change_direction(svc_pct)}{abs_pct_str(svc_pct)}。"
    )

    platform_note = build_platform_decline_text(cur_dau, prev_dau)
    trend_text    = analyze_dau_trend(cur_daily_dau)
    chart_path    = generate_dau_chart(
        cur_daily_dau, prev_daily_dau,
        cur_start, cur_end, prev_start, prev_end
    )

    def svc_change_name_str(lst):
        if not lst:
            return "暂无"
        parts = []
        for n, c, p, pct in lst:
            if pct == float('inf'):
                parts.append(f"{n}（新增服务）")
            else:
                parts.append(f"{n}（{abs_pct_str(pct)}）")
        return "、".join(parts)

    # ── 生成 Word 文档 ────────────────────
    doc = Document()
    for section in doc.sections:
        section.top_margin    = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin   = Inches(1.25)
        section.right_margin  = Inches(1.25)

    # 标题
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_after       = Pt(4)
    title_p.paragraph_format.first_line_indent = Pt(0)
    set_font(title_p.add_run("日活、服务、总结"), bold=True, size=16)

    note_p = doc.add_paragraph()
    note_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    note_p.paragraph_format.space_after       = Pt(8)
    note_p.paragraph_format.first_line_indent = Pt(0)
    set_font(note_p.add_run("默认日活含社区保障e管家"),
             size=10.5, color=RGBColor(0x60, 0x60, 0x60))

    # ══════════════════════════════════════
    #  一、日活情况
    # ══════════════════════════════════════
    add_heading_para(doc, "一、日活情况")

    # 1. 日活整体数据
    add_para(doc, "1.日活整体数据：", bold=True, first_line_indent=False)
    add_para(doc, dau_text_main)
    add_blank_para(doc, "（主要原因分析，请人工填写）")

    # 各端日活变化备注（①优化点：下降时列出下降端详情）
    add_para(doc, platform_note)

    # 2. 日活本周期趋势（②优化点：自动描述走势 + 折线图）
    add_para(doc, "2.日活本周期趋势：", bold=True, first_line_indent=False)
    add_para(doc, trend_text)

    # 插入折线图
    pic_para = doc.add_paragraph()
    pic_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pic_para.paragraph_format.first_line_indent = Pt(0)
    pic_para.paragraph_format.space_after       = Pt(6)
    run = pic_para.add_run()
    run.add_picture(chart_path, width=Inches(6.5))

    # ══════════════════════════════════════
    #  二、服务人次
    # ══════════════════════════════════════
    add_heading_para(doc, "二、服务人次")

    add_para(doc, "1.服务人次整体数据：", bold=True, first_line_indent=False)
    add_para(doc, svc_text_main)

    add_para(doc, "2.服务人次事项涨跌榜：本周或上周总使用人次1000以上的用户中",
             bold=True, first_line_indent=False)

    # 涨幅前三
    if up_list:
        add_para(doc, f"涨幅前三：{svc_change_name_str(up_list)}")
        conn2 = pymysql.connect(**DB_CONFIG)
        c2    = conn2.cursor()
        for idx, (name, cur, prev, pct) in enumerate(up_list, 1):
            peak_d, peak_v = query_service_peak(c2, name, cur_start, cur_end)
            if pct == float('inf'):
                desc = f"本周期{cur:,}人次（上周期无数据）"
            else:
                desc = (
                    f"从上周期的{prev:,}人次上升至本周期{cur:,}人次，"
                    f"提升{abs_pct_str(pct)}"
                )
                if peak_d:
                    desc += f"。峰值为{fmt_date_cn(peak_d)}单日点击{peak_v:,}次"
            add_para(doc, f"{idx}）{name}：{desc}。")
        c2.close()
        conn2.close()
    else:
        add_para(doc, "涨幅前三：本周期无明显上升的服务。")

    add_para(doc, "其他需要注意上升服务/功能：暂无。")

    # 跌幅前三
    if down_list:
        add_para(doc, f"跌幅前三：{svc_change_name_str(down_list)}")
        conn2 = pymysql.connect(**DB_CONFIG)
        c2    = conn2.cursor()
        for idx, (name, cur, prev, pct) in enumerate(down_list, 1):
            peak_d, peak_v = query_service_peak(c2, name, prev_start, prev_end)
            desc = (
                f"从上周期的{prev:,}人次下降至本周期{cur:,}人次，"
                f"下降{abs_pct_str(pct)}"
            )
            if peak_d:
                desc += f"。上周期峰值为{fmt_date_cn(peak_d)}单日{peak_v:,}次"
            add_para(doc, f"{idx}）{name}：{desc}。")
        c2.close()
        conn2.close()
    else:
        add_para(doc, "跌幅前三：本周期无下跌的服务。")

    # ══════════════════════════════════════
    #  三、本周总结
    # ══════════════════════════════════════
    add_heading_para(doc, "三、本周总结")

    add_para(doc, f"1）日活趋势：{dau_text_main}")
    add_blank_para(doc, "（主要原因分析，请人工填写）")

    add_para(doc, "2）人均行为次数：")
    add_blank_para(doc, "（人均行为次数分析，请人工填写）")

    add_para(doc, "3）本周宣推事件：")
    add_blank_para(doc, "（本周宣推事件，请人工填写）")

    add_para(doc, "4）搜索词变化：")
    if search_ok:
        add_search_table(doc, top_delta, top_cur)
        doc.add_paragraph()
        add_blank_para(doc, "（搜索词变化分析，请人工填写）")
    else:
        add_blank_para(
            doc,
            "（搜索词数据尚未导入，请先运行 search_detail_import.py 后重新生成报告）"
        )

    # 保存
    year_short  = str(cur_start.year)[2:]
    filename    = (
        f"周报-{year_short}年{cur_start.month}月{cur_start.day}日"
        f"-{cur_end.month}月{cur_end.day}日.docx"
    )
    output_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), filename)
    doc.save(output_path)
    print(f"\n[OK] 周报已生成：{output_path}")
    return output_path


if __name__ == "__main__":
    generate_report()

# -*- coding: utf-8 -*-
"""
report_engine.py — 通用报告生成引擎

用法：
    python report_engine.py --template templates/my_report.docx --config templates/my_report.json

模板目录约定（--template-dir 指定，默认 ./templates）：
    templates/
        my_report.docx   ← Word 模板，用 {{变量名}} 标记占位符
        my_report.json   ← 字段配置（与模板同名，扩展名换 .json）

也可用 --template / --config 分别指定路径。

配置文件格式见 README.md 或 templates/example.json。
"""

import argparse
import datetime
import json
import os
import re
import sys

import pymysql
from docx import Document
from docx.oxml.ns import qn

# ── 环境变量支持 ──────────────────────────────────────────────
try:
    from dotenv import load_dotenv
    load_dotenv()
except ImportError:
    pass

DB_CONFIG = {
    'host':     os.getenv('DB_HOST',     'localhost'),
    'port':     int(os.getenv('DB_PORT', '3306')),
    'user':     os.getenv('DB_USER',     'root'),
    'password': os.getenv('DB_PASSWORD', ''),
    'database': os.getenv('DB_NAME',     'daily'),
    'charset':  'utf8mb4',
}


# ── 内置日期变量 ──────────────────────────────────────────────
def _builtin_dates():
    today     = datetime.date.today()
    yesterday = today - datetime.timedelta(days=1)
    last_week = yesterday - datetime.timedelta(days=7)
    return {
        'today':     str(today),
        'yesterday': str(yesterday),
        'last_week': str(last_week),
        'date':      datetime.datetime.now().strftime('%Y%m%d_%H%M%S'),
    }


# ── 格式化函数 ────────────────────────────────────────────────
def _fmt(value, fmt: str) -> str:
    if value is None:
        return ''
    try:
        v = float(value)
    except (TypeError, ValueError):
        return str(value)

    if fmt == 'wan':
        return f'{v / 10000:.2f}万'
    if fmt == 'pct':
        sign = '+' if v >= 0 else ''
        return f'{sign}{v * 100:.2f}%'
    if fmt == 'int':
        return f'{int(v):,}'
    if fmt == 'date_cn':
        # 支持 date 对象或 YYYY-MM-DD 字符串
        try:
            d = datetime.date.fromisoformat(str(value)[:10])
            return f'{d.month}月{d.day}日'
        except Exception:
            return str(value)
    return str(value)


# ── SQL 查询 ──────────────────────────────────────────────────
def _run_sql(query: str, conn) -> object:
    """执行 SQL，返回第一行第一列的值；无结果返回 None。"""
    with conn.cursor() as cur:
        cur.execute(query)
        row = cur.fetchone()
        return row[0] if row else None


# ── 变量解析 ──────────────────────────────────────────────────
def resolve_variables(config: dict, conn) -> dict:
    """
    根据配置解析所有变量，返回 {变量名: 格式化后的字符串}。
    """
    dates = _builtin_dates()
    result = dict(dates)  # 内置日期变量直接可用

    for name, spec in config.get('variables', {}).items():
        vtype = spec.get('type', 'raw')
        fmt   = spec.get('format', 'raw')

        if vtype == 'date':
            val = spec.get('value', 'today')
            raw = dates.get(val, val)
        elif vtype == 'sql':
            raw_query = spec['query']
            # 将内置日期变量替换进 SQL
            for k, v in dates.items():
                raw_query = raw_query.replace(f'{{{k}}}', v)
            try:
                raw = _run_sql(raw_query, conn)
            except Exception as e:
                print(f'  [WARN] 变量 {name} 查询失败: {e}')
                raw = None
        elif vtype == 'literal':
            raw = spec.get('value', '')
        else:
            raw = spec.get('value', '')

        result[name] = _fmt(raw, fmt)

    return result


# ── Word 模板替换 ─────────────────────────────────────────────
_PLACEHOLDER_RE = re.compile(r'\{\{(\w+)\}\}')


def _replace_in_run(run, variables: dict):
    text = run.text
    def replacer(m):
        key = m.group(1)
        return variables.get(key, m.group(0))
    run.text = _PLACEHOLDER_RE.sub(replacer, text)


def _replace_in_paragraph(para, variables: dict):
    # 先尝试逐 run 替换（保留格式）
    for run in para.runs:
        _replace_in_run(run, variables)

    # 如果占位符跨 run 被拆散，做全文合并替换（会丢失部分格式，但保证内容正确）
    full = para.text
    if _PLACEHOLDER_RE.search(full):
        def replacer(m):
            return variables.get(m.group(1), m.group(0))
        new_text = _PLACEHOLDER_RE.sub(replacer, full)
        if new_text != full:
            # 清空所有 run，写入第一个 run
            for i, run in enumerate(para.runs):
                run.text = new_text if i == 0 else ''


def fill_template(template_path: str, variables: dict, output_path: str):
    doc = Document(template_path)

    # 正文段落
    for para in doc.paragraphs:
        _replace_in_paragraph(para, variables)

    # 表格单元格
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    _replace_in_paragraph(para, variables)

    # 页眉页脚
    for section in doc.sections:
        for hdr in (section.header, section.footer):
            if hdr:
                for para in hdr.paragraphs:
                    _replace_in_paragraph(para, variables)

    os.makedirs(os.path.dirname(os.path.abspath(output_path)), exist_ok=True)
    doc.save(output_path)
    print(f'[OK] 报告已生成：{output_path}')


# ── 主流程 ────────────────────────────────────────────────────
def generate(template_path: str, config_path: str):
    with open(config_path, encoding='utf-8') as f:
        config = json.load(f)

    report_name = config.get('report_name', '报告')
    output_dir  = config.get('output_dir', '报表产出')
    filename_tpl = config.get('output_filename', f'{report_name}_{{date}}.docx')

    print(f'正在生成：{report_name}')
    print(f'  模板：{template_path}')
    print(f'  配置：{config_path}')

    conn = pymysql.connect(**DB_CONFIG)
    try:
        variables = resolve_variables(config, conn)
    finally:
        conn.close()

    # 输出文件名中也支持变量替换
    filename = filename_tpl
    for k, v in variables.items():
        filename = filename.replace(f'{{{k}}}', v)

    output_path = os.path.join(output_dir, filename)
    fill_template(template_path, variables, output_path)
    return output_path


# ── CLI ───────────────────────────────────────────────────────
def _find_config(template_path: str) -> str:
    """同名 .json 配置文件自动发现。"""
    base = os.path.splitext(template_path)[0]
    for ext in ('.json', '.JSON'):
        p = base + ext
        if os.path.exists(p):
            return p
    return ''


def main():
    parser = argparse.ArgumentParser(description='通用报告生成引擎')
    parser.add_argument('--template',     help='Word 模板路径（.docx）')
    parser.add_argument('--config',       help='字段配置路径（.json）')
    parser.add_argument('--template-dir', default='templates',
                        help='模板目录（默认 ./templates），自动扫描所有 .docx+.json 对')
    parser.add_argument('--list',         action='store_true',
                        help='列出模板目录中可用的报告模板')
    args = parser.parse_args()

    if args.list:
        tdir = args.template_dir
        if not os.path.isdir(tdir):
            print(f'模板目录不存在：{tdir}')
            sys.exit(1)
        pairs = []
        for f in os.listdir(tdir):
            if f.lower().endswith('.docx'):
                cfg = _find_config(os.path.join(tdir, f))
                pairs.append((f, '✓' if cfg else '✗ (缺少 .json)'))
        if not pairs:
            print(f'模板目录 {tdir} 中没有找到 .docx 文件')
        else:
            print(f'模板目录：{tdir}')
            for name, status in pairs:
                print(f'  {status}  {name}')
        return

    if args.template:
        template_path = args.template
        config_path   = args.config or _find_config(template_path)
        if not config_path:
            print(f'未找到配置文件，请用 --config 指定，或在模板同目录放置同名 .json')
            sys.exit(1)
        generate(template_path, config_path)
    else:
        # 批量处理模板目录
        tdir = args.template_dir
        if not os.path.isdir(tdir):
            print(f'模板目录不存在：{tdir}，请用 --template 指定单个模板')
            sys.exit(1)
        found = False
        for f in sorted(os.listdir(tdir)):
            if not f.lower().endswith('.docx'):
                continue
            tpl = os.path.join(tdir, f)
            cfg = _find_config(tpl)
            if not cfg:
                print(f'[SKIP] {f}（缺少同名 .json 配置）')
                continue
            found = True
            generate(tpl, cfg)
        if not found:
            print('没有找到可处理的模板，请检查模板目录或使用 --template 指定')


if __name__ == '__main__':
    main()

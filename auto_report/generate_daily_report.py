# -*- coding: utf-8 -*-
"""
generate_daily_report.py
天府市民云工作日报自动化生成脚本

文档格式（严格对应已生成样例）：
  - 标题+副标题：同一段落居中，方正小标宋_GBK 16pt / 楷体 14pt，用 <w:br/> 换行
  - 正文两段：仿宋 16pt，首行缩进 640 twips（=两字符）
"""

import os
import datetime
import pandas as pd
import matplotlib.pyplot as plt
import matplotlib
from docx import Document
from docx.shared import Pt, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH
from sqlalchemy import create_engine

# 设置 matplotlib 中文字体
matplotlib.rcParams['font.sans-serif'] = ['SimHei', 'Microsoft YaHei']
matplotlib.rcParams['axes.unicode_minus'] = False

# ── 1. 配置 ──────────────────────────────────────────────────
# 数据库连接：优先读取环境变量，也可在此处直接填写
try:
    from dotenv import load_dotenv
    load_dotenv()
except ImportError:
    pass

DB_CONFIG = {
    'host':     os.getenv('DB_HOST',     'localhost'),
    'port':     int(os.getenv('DB_PORT', '3306')),
    'user':     os.getenv('DB_USER',     'root'),
    'password': os.getenv('DB_PASSWORD', ''),
    'database': os.getenv('DB_NAME',     'daily'),
    'charset':  'utf8mb4'
}

# 路径：优先读取环境变量，默认指向本脚本同目录下的子目录
SERVICE_MAPPING_PATH = os.getenv(
    'SERVICE_MAPPING_PATH',
    os.path.join(os.path.dirname(os.path.abspath(__file__)), '历史数据', '是否为服务.xlsx')
)
OUTPUT_DIR = os.getenv(
    'OUTPUT_DIR',
    os.path.join(os.path.dirname(os.path.abspath(__file__)), '报表产出')
)
os.makedirs(OUTPUT_DIR, exist_ok=True)

engine = create_engine(
    f"mysql+pymysql://{DB_CONFIG['user']}:{DB_CONFIG['password']}"
    f"@{DB_CONFIG['host']}:{DB_CONFIG['port']}/{DB_CONFIG['database']}?charset=utf8mb4"
)

YESTERDAY = datetime.date.today() - datetime.timedelta(days=1)
LAST_WEEK  = YESTERDAY - datetime.timedelta(days=7)


# ── 2. 辅助函数 ───────────────────────────────────────────────

def to_wan(val):
    """转换为万，保留两位小数"""
    return f"{val / 10000:.2f}"


def to_pct(val):
    """转为百分比字符串，正数加 + 号"""
    return f"{'+' if val >= 0 else ''}{val * 100:.2f}%"


def set_run_font(run, font_name, size_pt):
    """统一设置 run 的中英文字体和字号"""
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn('w:ascii'),   font_name)
    rFonts.set(qn('w:hAnsi'),   font_name)
    rFonts.set(qn('w:eastAsia'), font_name)
    run.font.size = Pt(size_pt)


def add_first_line_indent(paragraph, twips=640):
    """给段落添加首行缩进（单位 twips）"""
    pPr = paragraph._element.get_or_add_pPr()
    ind = OxmlElement('w:ind')
    ind.set(qn('w:firstLine'), str(twips))
    pPr.append(ind)


def generate_chart(data, output_path):
    """生成增幅前5服务柱形图（仅昨日增幅，按降序排列）"""
    top5 = data['top5']

    names = [name[:8] for name in top5.index]
    y_vals = top5['yesterday'].values
    pct_vals = top5['pct_change'].values

    fig, ax = plt.subplots(figsize=(8, 5.5))

    x = range(len(names))
    bars = ax.bar(x, pct_vals * 100, color='#4472C4', alpha=0.9, width=0.5)

    ax.set_xlabel('服务名称', fontsize=11)
    ax.set_ylabel('增幅（%）', fontsize=11)
    ax.set_title('昨日增幅前5位服务（周同比）', fontsize=13, fontweight='bold', pad=15)
    ax.set_xticks(x)
    ax.set_xticklabels(names, fontsize=10)
    ax.grid(axis='y', alpha=0.3, linestyle='--')

    # 柱子上标注：增幅 + 昨日使用次数
    for bar, pct, val in zip(bars, pct_vals, y_vals):
        height = bar.get_height()
        label_text = f"{pct*100:+.1f}%\n{int(val):,}次"
        ax.text(bar.get_x() + bar.get_width()/2, height,
                label_text,
                ha='center', va='bottom', fontsize=9,
                fontweight='bold', color='#2E5090')

    plt.tight_layout()
    plt.savefig(output_path, dpi=200, bbox_inches='tight', facecolor='white')
    plt.close()
    print(f"  图表已生成: {output_path}")


# ── 3. 数据处理 ──────────────────────────────────────────────

def get_report_data():
    # A. 平台汇总指标（昨日 & 上周同期）
    df_metrics = pd.read_sql(
        "SELECT * FROM platform_daily_metrics WHERE stat_date IN (%s, %s)",
        engine, params=(YESTERDAY, LAST_WEEK)
    )
    df_metrics['stat_date'] = pd.to_datetime(df_metrics['stat_date']).dt.date
    df_metrics.set_index('stat_date', inplace=True)

    # B. 核心服务名单（来自 Excel 映射表）
    df_map = pd.read_excel(SERVICE_MAPPING_PATH)
    df_map.columns = [c.strip() for c in df_map.columns]
    core_list = set(
        str(x).strip()
        for x in df_map[df_map['是否为服务'] == 1]['服务名称'].tolist()
    )

    # C. 服务明细（多端口合并后，仅取昨日&上周）
    df_detail = pd.read_sql(
        "SELECT DATE(stat_date) AS stat_date, service_name, "
        "       SUM(service_amount) AS service_amount "
        "FROM `5100_detail` "
        "WHERE DATE(stat_date) IN (%s, %s) "
        "GROUP BY DATE(stat_date), service_name",
        engine, params=(YESTERDAY, LAST_WEEK)
    )
    df_detail['stat_date'] = pd.to_datetime(df_detail['stat_date']).dt.date

    # D. 核心指标
    y_dau      = int(df_metrics.loc[YESTERDAY, 'platform_dau'])       if YESTERDAY in df_metrics.index else 0
    lw_dau     = int(df_metrics.loc[LAST_WEEK,  'platform_dau'])       if LAST_WEEK  in df_metrics.index else 0
    dau_growth = (y_dau - lw_dau) / lw_dau if lw_dau > 0 else 0

    # 核心功能服务人次（仅统计"是否为服务=1"的服务）
    y_core_sum = int(
        df_detail[
            (df_detail['stat_date'] == YESTERDAY) &
            (df_detail['service_name'].isin(core_list))  # core_list 来自 是否为服务=1
        ]['service_amount'].sum()
    )
    y_new_reg   = int(df_metrics.loc[YESTERDAY, 'new_register_users'])   if YESTERDAY in df_metrics.index else 0
    y_total_reg = int(df_metrics.loc[YESTERDAY, 'total_register_users']) if YESTERDAY in df_metrics.index else 0

    # E. 增幅前5服务（仅限核心服务，昨日 > 100 且上周有数据）
    y_grp  = df_detail[df_detail['stat_date'] == YESTERDAY].set_index('service_name')['service_amount']
    lw_grp = df_detail[df_detail['stat_date'] == LAST_WEEK ].set_index('service_name')['service_amount']

    df_growth = pd.DataFrame({'yesterday': y_grp, 'lastweek': lw_grp}).fillna(0)
    # 筛选：必须是核心服务 + 昨日>100 + 上周有数据
    df_growth = df_growth[
        df_growth.index.isin(core_list) &
        (df_growth['yesterday'] > 100) &
        (df_growth['lastweek'] > 0)
    ]
    df_growth['pct_change'] = (df_growth['yesterday'] - df_growth['lastweek']) / df_growth['lastweek']
    top5 = df_growth.sort_values('pct_change', ascending=False).head(5)

    return {
        'y_dau':      y_dau,
        'dau_growth': dau_growth,
        'y_core_sum': y_core_sum,
        'y_new_reg':  y_new_reg,
        'y_total_reg': y_total_reg,
        'top5':       top5,
    }


# ── 4. Word 文档生成 ─────────────────────────────────────────

def build_document(data):
    doc = Document()

    # 清除 python-docx 自动添加的空段落
    for p in list(doc.paragraphs):
        p._element.getparent().remove(p._element)

    # ---- 标题 + 副标题（同一段落，居中） ----
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 标题文字
    r_title = p_title.add_run(f"天府市民云{YESTERDAY.month}月{YESTERDAY.day}日运行情况")
    set_run_font(r_title, '方正小标宋_GBK', 16)

    # <w:br/> 换行（不分段）
    r_title._element.append(OxmlElement('w:br'))

    # 副标题文字
    r_sub = p_title.add_run("（报市城运中心当日值班组）")
    set_run_font(r_sub, '楷体', 14)

    # ---- 正文第一段 ----
    p1 = doc.add_paragraph()
    add_first_line_indent(p1, 640)

    text1 = (
        f"{YESTERDAY.year}年{YESTERDAY.month}月{YESTERDAY.day}日，"
        f"当日活跃用户{to_wan(data['y_dau'])}万，增幅{to_pct(data['dau_growth'])}，"
        f"当日核心功能服务{to_wan(data['y_core_sum'])}万人次，"
        f"当日新增注册用户数{data['y_new_reg']}人，"
        f"累计注册总用户数{to_wan(data['y_total_reg'])}万人。"
    )
    r1 = p1.add_run(text1)
    set_run_font(r1, '仿宋', 16)

    # ---- 正文第二段 ----
    p2 = doc.add_paragraph()
    add_first_line_indent(p2, 640)

    top5_str = "、".join(
        f"{name}（{to_pct(row['pct_change'])}）"
        for name, row in data['top5'].iterrows()
    )
    text2 = f"从服务项目使用情况：当日使用增幅排前5位的是{top5_str}。"
    r2 = p2.add_run(text2)
    set_run_font(r2, '仿宋', 16)

    # ---- 图表 ----
    chart_path = os.path.join(OUTPUT_DIR, '_chart_tmp.png')
    generate_chart(data, chart_path)
    doc.add_picture(chart_path, width=Inches(6))
    os.remove(chart_path)

    return doc


# ── 5. 主流程 ────────────────────────────────────────────────

def main():
    print(f"正在生成 {YESTERDAY} 工作日报...")

    try:
        data = get_report_data()
    except Exception as e:
        print(f"数据获取失败: {e}")
        raise

    print(f"  平台DAU:   {to_wan(data['y_dau'])}万  周同比增幅: {to_pct(data['dau_growth'])}")
    print(f"  核心服务:  {to_wan(data['y_core_sum'])}万人次")
    print(f"  新增注册:  {data['y_new_reg']}人    累计注册: {to_wan(data['y_total_reg'])}万人")
    print(f"  增幅前5:   {list(data['top5'].index)}")

    doc = build_document(data)

    fname     = f"天府市民云工作日报-城运中心{YESTERDAY.strftime('%Y%m%d')}.docx"
    save_path = os.path.join(OUTPUT_DIR, fname)
    doc.save(save_path)
    print(f"报表已保存: {save_path}")


if __name__ == "__main__":
    main()
